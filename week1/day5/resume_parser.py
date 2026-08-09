import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API key kaha hai bhai")

client=Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"

 # Part 1: Job Description Parsing
job_description="""
Description
# Job Title: AI/ML Engineer (Fresher)

## Location

Bengaluru, India (Hybrid/On-site)

## Job Type

Full-time

## About the Role

We are looking for a passionate and motivated AI/ML Engineer (Fresher) to join our team. This role is ideal for recent graduates or entry-level candidates who are eager to build intelligent applications using Machine Learning, Deep Learning, and Generative AI technologies. You will work alongside experienced engineers to develop, test, and deploy AI-powered solutions for real-world business problems.

## Key Responsibilities

* Develop and implement machine learning models for predictive analytics and automation.
* Collect, clean, and preprocess structured and unstructured datasets.
* Perform exploratory data analysis (EDA) to identify insights and trends.
* Build AI applications using Python and popular ML libraries.
* Assist in training, evaluating, and optimizing machine learning models.
* Work with APIs and integrate AI models into applications.
* Collaborate with software developers, data engineers, and product teams.
* Document experiments, model performance, and technical findings.
* Stay updated with the latest advancements in AI, Machine Learning, and Generative AI.

## Required Skills

* Strong programming skills in Python.
* Basic understanding of Machine Learning algorithms.
* Knowledge of data structures and algorithms.
* Familiarity with NumPy, Pandas, Matplotlib, and Scikit-learn.
* Basic understanding of Deep Learning concepts.
* Knowledge of SQL and database fundamentals.
* Understanding of statistics, probability, and linear algebra.
* Good analytical and problem-solving skills.
* Excellent communication and teamwork abilities.

## Preferred Skills

* Experience with TensorFlow or PyTorch.
* Knowledge of Large Language Models (LLMs) and Prompt Engineering.
* Familiarity with LangChain or LlamaIndex.
* Understanding of Retrieval-Augmented Generation (RAG).
* Exposure to vector databases such as FAISS, Chroma, or Pinecone.
* Basic knowledge of cloud platforms (AWS, Azure, or Google Cloud).
* Experience with Git and GitHub.

## Educational Qualification

* Bachelor's degree in Computer Science, Artificial Intelligence, Information Technology, Data Science, Electronics, or a related STEM field.
* Candidates graduating in the current or previous academic year are encouraged to apply.

## Nice-to-Have Projects

* Chatbots using LLMs
* Recommendation Systems
* Image Classification
* Sentiment Analysis
* Resume Parser
* Question Answering Systems
* RAG-based Applications
* AI Automation Projects

## What We Offer

* Mentorship from experienced AI professionals.
* Opportunity to work on real-world AI and Generative AI projects.
* Hands-on experience with modern AI frameworks and tools.
* Learning and certification support.
* Collaborative and innovation-driven work environment.
* Competitive salary and career growth opportunities.

## Selection Process

1. Online Aptitude Assessment
2. Python Programming Test
3. Machine Learning Fundamentals Interview
4. Technical Interview
5. HR Interview

"""
class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are an expert HR assistant.

Your job is to analyze job descriptions and extract
structured information from them.

Return ONLY valid JSON matching this schema:

{jobd_schema}
IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""

user_prompt = f"""
Analyze the following job description:

{job_description}
"""
message_system={
    "role" : "system",
    "content" : system_prompt
}
message_user={
    "role" : "user",
    "content" : user_prompt
}
response_format={
    "type" : "json_object"
}


messages=[message_system, message_user]

response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)


answer=response.choices[0].message.content

raw_json=answer
# print(raw_json)



import json
job_data=json.loads(raw_json)

job = JobD(**job_data)

print(job.minimum_experience)
print(job.education_requirements)



#parse real
#class MatchResult(BaseModel):
   # score: float
    #details: dict #details means what is missing and what is strength of the candidate

class MatchResult(BaseModel):
    candidate_name: str

    overall_score: float

    skill_score: float
    experience_score: float
    education_score: float
    project_score: float

    matching_skills: list[str]
    missing_skills: list[str]

    strengths: list[str]
    weaknesses: list[str]

    ats_improvement_tips: list[str]

    learning_resources: list[str]

    estimated_score_after_improvement: float

    verdict: str


#Part 2: Resume Parsing
class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


resume_schema = Resume.model_json_schema() #Resume schema is created here

def final_score(job,resume):
    match_schema = MatchResult.model_json_schema() #first create schema of match result and then pass it to the llm in prompt
    prompt = f"""
You are an expert ATS (Applicant Tracking System) recruiter with 15+ years of experience in technical hiring.

Your task is to compare the candidate's resume with the job description and return ONLY valid JSON matching the schema below.

JOB DESCRIPTION:
{job.model_dump_json(indent=2)}

CANDIDATE RESUME:
{resume.model_dump_json(indent=2)}

JSON SCHEMA:
{match_schema}

Evaluation Instructions:

Evaluate the candidate using the following weights:

• Skills Match → 40%
• Experience → 25%
• Projects → 20%
• Education → 15%

Carefully compare:

- Required Skills
- Preferred Skills
- Relevant Experience
- Projects
- Certifications
- Education

Return ONLY the following fields exactly as defined in the schema.

candidate_name
- Candidate's full name.

overall_score
- Overall ATS score from 0 to 100.

skill_score
- Score based only on technical skills.

experience_score
- Score based only on work experience and internships.

education_score
- Score based on education.

project_score
- Score based on relevance and quality of projects.

matching_skills
- List every matching technical skill.

missing_skills
- List important missing skills.

strengths
- 3–5 concise strengths.

weaknesses
- 3–5 concise weaknesses.

ats_improvement_tips

Give exactly 5 practical ATS improvement suggestions.

learning_resources

Recommend one learning resource for each missing skill.

Examples

TensorFlow → TensorFlow Official Tutorials

Docker → Docker Docs

LangChain → LangChain Documentation

FAISS → FAISS Documentation

estimated_score_after_improvement

Estimate the ATS score after the candidate completes all suggested improvements.

- Suggestions should be actionable.
- Example:
  - Learn TensorFlow
  - Build one RAG project
  - Add Docker experience
  - Add GitHub portfolio
  - Quantify project achievements

verdict
- One short hiring recommendation.
Examples:
"Excellent Match"
"Good Match"
"Needs Improvement"
"Not Recommended"

Rules:

1. Return ONLY valid JSON.
2. Do NOT return markdown.
3. Do NOT explain your reasoning.
4. Do NOT include extra keys.
5. Do NOT invent skills or experience.
6. If information is unavailable, use null or an empty list.
7. Keep all text concise and professional.
"""
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)

def parse_resume(resume_text):  #Function for parsing the resume text and returning a Resume object
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


#Part 3: Resume File Reading
from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):  #This func will return all the text from all pages of pdf file
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    #There can be tables in resume so we extract that text data 
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":  #If the file is pdf then we call read_pdf func to extract text from pdf
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx": #If the file is docx then we call read_docx func to extract text from docx
        return read_docx(file_path)
    else:
        return None

def print_ats_report(result):
    print("\n" + "=" * 60)
    print("           ATS RESUME ANALYSIS REPORT")
    print("=" * 60)

    print(f"\nCandidate : {result.candidate_name}")

    print(f"\nOverall ATS Score : {result.overall_score:.1f}%")

    print("\nDetailed Scores")
    print("-" * 40)
    print(f"Skills       : {result.skill_score:.1f}%")
    print(f"Experience   : {result.experience_score:.1f}%")
    print(f"Projects     : {result.project_score:.1f}%")
    print(f"Education    : {result.education_score:.1f}%")

    print("\nMatching Skills")
    print("-" * 40)
    for skill in result.matching_skills:
        print(f"✓ {skill}")

    print("\nMissing Skills")
    print("-" * 40)
    for skill in result.missing_skills:
        print(f"✗ {skill}")

    print("\nStrengths")
    print("-" * 40)
    for item in result.strengths:
        print(f"• {item}")

    print("\nWeaknesses")
    print("-" * 40)
    for item in result.weaknesses:
        print(f"• {item}")

    print("\nATS Improvement Tips")
    print("-" * 40)
    for tip in result.ats_improvement_tips:
        print(f"✓ {tip}")

    print("\nRecommended Learning Resources")
    print("-" * 40)
    for resource in result.learning_resources:
        print(f"📘 {resource}")

    print("\nEstimated ATS Score After Improvement")
    print("-" * 40)
    print(f"{result.estimated_score_after_improvement:.1f}%")

    print("\nVerdict")
    print("-" * 40)
    print(result.verdict)

    print("=" * 60)


#Now readding Resumes from the folder and parsing them and scoring them based on the job description
# lets do it now
resume_folder = Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():
    #C:\Users\krish\OneDrive\Desktop\AI Engineer\week1\day5\resumes\ResumeV2.pdf
    if file_path.suffix.lower() not in [".pdf", ".docx"]: #If the file is not pdf or docx then we skip it and continue to the next file
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path) #Whole resume text is extracted from the file and stored in resume_text variable as string
    parsed_resume=parse_resume(resume_text) # llm call1 this will give us the parsed resume in structured format as Resume object we craeted above
    time.sleep(10) #To prevent DDos Attack on the server we add a delay of 5 seconds before making the next request to the server. 
    #This is to prevent the server from getting overloaded with too many requests in a short period of time.
    # kisi ne chatgpt ka acc banaya aur request bhejna shhur krega millions
    #chattgot server jam ho jayega
    result = final_score(job, parsed_resume) #llm caLL2
    #Now in result we have = score and details
    time.sleep(10)
    print_ats_report(result)
    all_results.append({
    "name": result.candidate_name,
    "score": result.overall_score,
    "result": result
    })
all_results.sort(
    key=lambda candidate: candidate["score"],  #we are sorting the candidates based on their score in descending order so that the top candidates are at the top of the list
    reverse=True
)
top_2 = all_results[:2]  #Selecting Top 2 candidates from the sorted list of candidates based on their score
worst_2 = all_results[-2:] #Selecting Worst 2 candidates from the sorted list of candidates based on their score


print("TOP 2 CANDIDATES")
for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(f"Verdict : {candidate['result'].verdict}")
    print(f"Overall ATS Score : {candidate['score']}%")

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(f"Verdict : {candidate['result'].verdict}")
    print(f"Overall ATS Score : {candidate['score']}%")